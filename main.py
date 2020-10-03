from bs4 import BeautifulSoup
import grequests
import requests
import re
from docx import Document
import time

cfURL = "http://codeforces.com"
class Crawler:

    def getLink(self):
        content = requests.get(cfURL + "/problemset/page/1", params = {"order": "BY_SOLVED_DESC"}).text
        soup = BeautifulSoup(content, "lxml")
        problems = soup.find_all("a", href = re.compile("/problemset/problem/.*"))
        links = [cfURL + item.attrs['href'] for item in problems[::2]]
        return links

    def process_text(self, content):
        statement = BeautifulSoup(content, "lxml").find("div", class_ = "problem-statement")
        title = re.match("\w*\. (.*)", statement.find("div", class_ = "title").text).group(1)
        passages = statement.find("div", class_ = None)
        passage = ''.join(item.text + "\r\n" for item in passages.find_all("p"))
        return (title, passage)

    def exception(self, request, exception):
        print("Problem: {}: {}".format(request.url, exception))

    def async_down(self, urls):
        reqs = (grequests.get(url) for url in urls)
        res = grequests.imap(reqs, size = 5, exception_handler = self.exception)
        doc = Document()
        for item in res:
            title, passage = self.process_text(item.text)
            index = re.match(".*problem/(.*)/(.*)", item.url)
            title = index.group(1) + index.group(2) + '. '+ title
            doc.add_heading(title, 1)
            doc.add_paragraph(passage)
            print(title, "successfully written.")
        doc.save("cfProblems.docx")

if __name__ == '__main__':
    t = time.time()
    c = Crawler()
    links = c.getLink()
    print(str(len(links)) + " problems found on homepage.")

    c.async_down(links[:30])
    print("Completed. Time used:", time.time() - t, "s")
